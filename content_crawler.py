from warnings import warn

import re
import os.path as osp
from io import BytesIO

import requests

from bs4 import BeautifulSoup
try:
    no_docx = None
    from docx import Document
    from docx.shared import Inches
except ModuleNotFoundError as no_docx:
    warn(f'docx module is not found. Make sure docx is installed')

invalid_char = r'<>:"/\|?*'

def _insert_img(img_url: str, doc) -> None:
    image_from_url = requests.get(img_url)
    io_url = BytesIO(image_from_url.content)
    doc.add_picture(io_url, Inches(6))

def crawl_content(url: str, output_dir: str, *, to_docx: bool = True) -> None:
    """
    Tries to crawl text and images between HTML `<h1>` and `<h3>` tag to\
    `.docx` or `.txt` file.\n

    ## Parameters:
    url: website to crawl from\n
    output_dir: the directory to save output_file to\n
    to_docx: whether to save as `.docx` or as `.txt`
    """

    if to_docx and no_docx: raise no_docx

    html = requests.get(url).text
    soup = BeautifulSoup(html, 'html.parser')
    all_tag = soup.find_all(re.compile(r'^(?:p$|h[1-6]|img)'))


    start, end = 0, len(all_tag)
    for i, tag in enumerate(all_tag):
        if re.findall(r'^<h1', str(tag)):
            start = i
            break
    for i, tag in enumerate(all_tag):
        if re.findall(r'^<h3', str(tag)):
            end = i - 1
            break
    all_tag = all_tag[start:end]


    if to_docx:
        doc_file = Document()
        for tag in all_tag:
            if re.findall(r'^<[ph]', str(tag)):
                doc_file.add_paragraph(tag.get_text())
            elif 'data-src' in tag.attrs: #imag
                _insert_img(tag['data-src'], doc_file)
    else:
        article = ''
        for tag in all_tag:
            if re.findall(r'^<[ph]', str(tag)):
                article += tag.get_text() + '\n'
        

    title = ''
    for char in soup.title.string:
        if not char in invalid_char: title += char
    save_dir = osp.join(output_dir, title + '.docx')

    if to_docx: doc_file.save(save_dir)
    else:
        with open(save_dir, 'w', encoding='utf-8') as txt_file:
            txt_file.write(article)